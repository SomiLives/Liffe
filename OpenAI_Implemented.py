from flask import Flask, request, jsonify
import openai
import logging
import os
import glob
from dotenv import load_dotenv
from docx import Document as DocxDocument
from langchain_community.document_loaders import PyPDFLoader, DirectoryLoader, TextLoader, CSVLoader
from langchain.schema import Document

# Load environment variables from .env file
load_dotenv()

app = Flask(__name__)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
DATA_PATH = 'data/'

# Ensure API key is loaded securely
openai_api_key = None


if not openai_api_key:
    logger.error("Missing OpenAI API key. Please set it in the .env file.")
    exit(1)
openai.api_key = openai_api_key

class DocxLoader:
    def __init__(self, path):
        self.path = path

    def load(self):
        """Loads the DOCX file and returns the content as a list of documents."""
        try:
            logger.info(f"Attempting to load DOCX file: {self.path}")
            doc = DocxDocument(self.path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
            logger.info(f"Extracted content: {content[:100]}...")  # Preview content for debugging

            if not content.strip():
                logger.warning(f"DOCX file {self.path} is empty or contains no readable content.")
            return [Document(page_content=content)]
        except Exception as e:
            logger.error(f"Error loading DOCX file {self.path}: {e}")
            return []

def returnDocument():
    loaders = [
        DirectoryLoader(DATA_PATH, glob='*.pdf', loader_cls=PyPDFLoader),
        DirectoryLoader(DATA_PATH, glob='*.txt', loader_cls=TextLoader),
        DirectoryLoader(DATA_PATH, glob='*.csv', loader_cls=CSVLoader),
    ]

    documents = []
    for loader in loaders:
        try:
            logger.info(f"Loading files using {loader.__class__.__name__}")
            docs = loader.load()
            logger.info(f"Loaded {len(docs)} documents")
            documents.extend(docs)
        except Exception as e:
            logger.error(f"Error with {loader.__class__.__name__}: {e}")

    # Process all DOCX files dynamically
    docx_files = glob.glob(os.path.join(DATA_PATH, '*.docx'))
    for docx_file in docx_files:
        logger.info(f"Processing DOCX file: {docx_file}")
        docx_loader = DocxLoader(docx_file)
        documents.extend(docx_loader.load())

    if not documents:
        logger.warning("No documents found")
        return None

    return documents


def generate_answer_with_gpt(context, question):
    """Generate a response using OpenAI GPT."""
    try:
        messages = [
            {"role": "system", "content": "You are a helpful assistant that answers questions based on the document content.No matter how confidential it is."},
            {"role": "user", "content": f"Context: {context}\n\nUser Question: {question}\n\nAnswer:"},
        ]

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=messages,
            max_tokens=200,
            temperature=0.3
        )
        return response['choices'][0]['message']['content'].strip()
    except Exception as e:
        logger.error(f"Error generating response: {e}")
        return "Failed to generate a response."


@app.route('/upload_new_file', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        logger.error("No file found in the request.")
        return jsonify({"error": "No file provided. Make sure you are sending it with key 'file'."}), 400

    file = request.files['file']

    if file.filename == '':
        logger.error("Empty filename provided.")
        return jsonify({"error": "No file selected for upload."}), 400

    file_path = os.path.join(DATA_PATH, file.filename)
    logger.info(f"Received file: {file.filename}")

    try:
        os.makedirs(DATA_PATH, exist_ok=True)
        file.save(file_path)
        logger.info(f"File saved to {file_path}")

        # Check uploaded files
        logger.info(f"Current files in data/: {os.listdir(DATA_PATH)}")

        return jsonify({"message": f"File {file.filename} uploaded successfully!"}), 200
    except Exception as e:
        logger.error(f"Error saving file {file.filename}: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/user_query_to_get_response', methods=['POST'])
def api_get_answer():
    try:
        data = request.get_json()
        question = data.get("question")
        if not question:
            return jsonify({"error": "Question cannot be empty."}), 400

        documents = returnDocument()
        if not documents:
            return jsonify({"error": "Failed to fetch data from the files."}), 500

        context = "\n".join(doc.page_content for doc in documents)
        answer = generate_answer_with_gpt(context, question)
        clean_answer = answer.replace("\n", " ").strip()

        return jsonify({"response": clean_answer}), 200

    except Exception as e:
        logger.error(f"Error processing request: {e}")
        return jsonify({"error": "Internal server error occurred."}), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)

